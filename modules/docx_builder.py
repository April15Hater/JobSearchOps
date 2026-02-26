"""
modules/docx_builder.py — Build .docx files for resume and cover letter export.

Two modes:
  1. Template mode: open a user-supplied .docx, find the placeholder line, replace it.
  2. Scratch mode: build a cleanly formatted document from scratch.

Placeholder strings:
  Resume:       [RESUME_CONTENT]
  Cover letter: [COVER_LETTER_CONTENT]
"""
from __future__ import annotations
import io
import logging
import os
import re
from datetime import date

logger = logging.getLogger(__name__)

# Section header patterns — lines that are all-caps (or common header names)
_SECTION_HEADER_RE = re.compile(
    r"^(SUMMARY|PROFESSIONAL SUMMARY|EXPERIENCE|WORK EXPERIENCE|SKILLS|"
    r"TECHNICAL SKILLS|EDUCATION|CERTIFICATIONS|AWARDS|PROJECTS|"
    r"PUBLICATIONS|VOLUNTEER|REFERENCES|[A-Z][A-Z\s&/\-]{3,})$"
)


def _is_section_header(line: str) -> bool:
    return bool(_SECTION_HEADER_RE.match(line.strip()))


# ---------------------------------------------------------------------------
# Resume
# ---------------------------------------------------------------------------

def build_resume_docx(resume_text: str, template_path: str = None) -> bytes:
    """
    Build a resume .docx from resume_text.
    If template_path points to a valid .docx, inject content there.
    Otherwise build from scratch.
    Returns raw bytes suitable for a Flask response.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if template_path and os.path.isfile(template_path):
        return _inject_into_template(resume_text, template_path, "[RESUME_CONTENT]")

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    lines = resume_text.splitlines()
    first_line = True

    for raw_line in lines:
        line = raw_line.rstrip()

        if first_line and line:
            # Treat the first non-empty line as the candidate's name
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)  # dark navy
            first_line = False
            continue

        if not line:
            # Blank line — add a small spacer paragraph
            p = doc.add_paragraph("")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            continue

        if _is_section_header(line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)  # blue
            # Underline the section header
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "1D4ED8")
            pBdr.append(bottom)
            pPr.append(pBdr)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(line)
            run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Cover letter
# ---------------------------------------------------------------------------

def build_cover_letter_docx(cover_letter_text: str, template_path: str = None) -> bytes:
    """
    Build a cover letter .docx from cover_letter_text.
    If template_path points to a valid .docx, inject content there.
    Otherwise build from scratch.
    Returns raw bytes suitable for a Flask response.
    """
    from docx import Document
    from docx.shared import Pt, Inches

    if template_path and os.path.isfile(template_path):
        return _inject_into_template(cover_letter_text, template_path, "[COVER_LETTER_CONTENT]")

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Date header
    p = doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    p.paragraph_format.space_after = Pt(12)
    for run in p.runs:
        run.font.size = Pt(11)

    # Body paragraphs
    for para in cover_letter_text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        p = doc.add_paragraph(para)
        p.paragraph_format.space_after = Pt(10)
        for run in p.runs:
            run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Template injection
# ---------------------------------------------------------------------------

def _inject_into_template(content: str, template_path: str, placeholder: str) -> bytes:
    """
    Open a .docx template, find the paragraph containing placeholder,
    replace it with the content lines, and return bytes.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document(template_path)

    target_para = None
    for para in doc.paragraphs:
        if placeholder in para.text:
            target_para = para
            break

    if target_para is None:
        logger.warning(
            "Placeholder '%s' not found in template %s — falling back to appending content.",
            placeholder, template_path
        )
        # Fall back: just append the content at the end
        for line in content.splitlines():
            p = doc.add_paragraph(line.rstrip())
            for run in p.runs:
                run.font.size = Pt(10)
    else:
        # Clear placeholder paragraph and insert content lines before it
        from docx.oxml.ns import qn
        parent = target_para._element.getparent()
        idx = list(parent).index(target_para._element)

        # Remove the placeholder paragraph
        parent.remove(target_para._element)

        # Insert new paragraphs in its place
        from docx.oxml import OxmlElement
        for i, line in enumerate(content.splitlines()):
            new_p = OxmlElement("w:p")
            new_r = OxmlElement("w:r")
            new_t = OxmlElement("w:t")
            new_t.text = line.rstrip()
            new_r.append(new_t)
            new_p.append(new_r)
            parent.insert(idx + i, new_p)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
